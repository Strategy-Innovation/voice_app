#!/usr/bin/env python
# coding: utf-8

# In[ ]:


from pydub import AudioSegment
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
import docx
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from docx import Document

from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

from flask import Flask, render_template, request, jsonify
import pandas as pd
from collections import Counter
import openai
import os
from flask_cors import CORS

app = Flask(__name__)
CORS(app)


# Set your OpenAI GPT-3 API key here
openai.api_key = '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'
#openai.api_key = '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'
client = openai.api_key

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Get form data
        print("RECORDINGGG!")
        audio_file = request.files['audio']
        audio_file_path = '/Users/strategy_innovation/Desktop/Risk Reco - App/Audio file/audio_mohit.m4a'
        audio_file.save(audio_file_path)
        m4a_file = '/Users/strategy_innovation/Desktop/Risk Reco - App/Audio file/audio_mohit.m4a' # I have downloaded sample audio from this link https://getsamplefiles.com/sample-audio-files/m4a
        wav_filename = '/Users/strategy_innovation/Desktop/Risk Reco - App/Audio file/audio_mohit.wav'

        sound = AudioSegment.from_file(m4a_file, format='m4a')
        file_handle = sound.export(wav_filename, format='wav')

        audio_file=open("/Users/strategy_innovation/Desktop/Risk Reco - App/Audio file/audio_mohit.wav","rb")
        transcript=openai.Audio.transcribe(api_key=openai.api_key, model="whisper-1",file=audio_file, language="en")
        print("Recording received")
        
        print(audio_file_path)
        text_1=transcript.text + "\n"
        print(text_1)

        # Define the prompt for risk recommendation
#         prompt = f"Given the provided transcript:\n\n{transcript}\n\nGenerate a comprehensive report by meticulously identifying the customer's name, Sum Insured, and the Industry type. If Industry type is not discernible, assign 'NA' accordingly. Present this information in a structured, row-wise format. Craft the report in a professional style, aligning with industry norms, and emphasize key industrial standards. Additionally, offer point-wise risk recommendations in accordance with environmental health and safety norms. Analyze the observations carefully and leverage OpenAI input to enhance the report, drawing insights from general knowledge."

#         prompt = (f"Given the provided transcript:\n\n{transcript}\n\n"
#     "Based on industry-specific observations, your task is to generate a detailed report. "
#     "Identify the customer's name, Sum Insured, and Industry type. If the Industry type is not available, assign 'NA'. "
#     "Present this information row-wise in a professional report format"
#     "Additionally, include point-wise risk recommendations aligned with environmental, health, industrial processes and safety norms related to that client specific industry type. "
#     "Analyze the recommendations and ensure that the recommendations are based on relevant industry saftey standards code and provide them in the recommendation report."
# )
        
#         prompt = (f"Given the provided transcript:\n\n{transcript}\n\n"
#     "From the the provided transcript, identify the customer's name, Sum Insured, and Industry type. If the Industry type is not available, assign 'NA'. "
#     "Then based on industry-specific observations, your task is to generate a proper detailed report in the following format sequentially and not in tabular format: Observation in sequential manner, Hazard, Recommendation, IS Reference code and it's respective clause governing the recommendation."
#     "The detailed report that you will generate should contain a sequential mapping of the observations recorded with the hazards that the recommended observation could pose to the industry and respective measures required to combat the problem inferred from the observation and recommendations based on your knowledge. The recommendations should be backed by Indian Safety standard codes as per the industry and hazard type and mention these respective safety standard codes that needs to be adhered in the report as well."
#     "Present this information row-wise in a professional report format."
#     "Additionally, include point-wise risk recommendations aligned with environmental, health, industrial processes and safety norms related to that client specific industry type. "
#     "Analyze the recommendations and ensure that the recommendations are based on relevant industry saftey standards code and provide them in the recommendation report."
# )
        prompt = (f"Given the provided transcript:\n\n{transcript}\n\n"
    "Based on industry-specific observations, your task is to generate a detailed report in the following format sequentially and not in tabular format: Observation in sequential manner, Hazard, Recommendation, IS code and it's respective clause governing the recommendation."
    "Using the transcript, you will get details like customer name, sum insured and industry type which needs to be shown in the top of the report.If the Industry type is not available, assign 'NA'. "
    "The detailed report that you will generate should contain a sequential mapping of the observations recorded with the hazards that the recommended observation could pose to the industry and respective measures required to combat the problem inferred from the observation and recommendations based on your knowledge. The recommendations should be backed by Indian Safety standard codes as per the industry and hazard type and mention these respective safety standard codes that needs to be adhered in the report as well."
    "Present this information row-wise in a professional report format"
    "Additionally, include point-wise risk recommendations aligned with environmental, health, industrial processes and safety norms related to that client specific industry type. "
    "Analyze the recommendations and ensure that the recommendations are based on relevant industry saftey standards code and provide them in the recommendation report."
    "Output the report containing the heading: Hazard and Recommendation in bold."
)

    
    
        # Call the OpenAI API to generate the risk recommendation
        response = openai.Completion.create(
         model="text-davinci-003",  # You can choose a different model if needed
         prompt=prompt,
         temperature=0.7,  # Adjust temperature for randomness
         max_tokens=800,  # Adjust max_tokens to control response length
         n=1  # Number of completions to generate
         )

        # Get the generated risk recommendation
        risk_recommendation = response.choices[0].text.strip()

        # Print the generated risk recommendation
        print("Generated Risk Recommendation:")
        print(risk_recommendation)
        
        
#         def save_to_word(risk_recommendation):s
#     # Create a new Word document
#             doc = Document()

#     # Add a heading to the document with bold and centrally aligned
#             heading = doc.add_heading('Generated Risk Recommendation', level=1)
#             heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#             heading.bold = True

#     # Add a line after the heading in brown color
#             line = heading.add_run()
#             line.add_break()
#             border = line.add_border_bottom()
#             border.color.rgb = RGBColor(165, 42, 42)  # Brown color

#     # Add the risk recommendation to the document
#             doc.add_paragraph(risk_recommendation)

#     # Set border for the entire document with blue color
#             section = doc.sections[0]
#             section_start = section.start_type
#             section_end = section.end_type

#             border_element = OxmlElement('w:sectBorder')
#             border_element.set(qn('w:val'), 'single')
#             border_element.set(qn('w:sz'), '4')
#             border_element.set(qn('w:color'), '0000FF')  # Blue color

#             section_start.insert_element_before(border_element)
#             section_end.insert_element_before(border_element)

#     # Save the document to a file
#             doc.save('risk_recommendation.docx')
#         # Assuming risk_recommendation is the generated recommendation
#             risk_recommendation = (risk_recommendation)

#         # Call the function to save to Word file
#             save_to_word(risk_recommendation)

#             print("Risk recommendation saved to 'risk_recommendation.docx'")
        
        # Function to save risk recommendation to Word file
        def save_to_word(risk_recommendation):
        # Create a new Word document
            doc = Document()

        # Add a heading to the document
            doc.add_heading('Generated Risk Recommendation', level=1)
#             heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#             heading.bold = True
 
        # Add the risk recommendation to the document
            doc.add_paragraph(risk_recommendation)

        # Save the document to a file
            doc.save('risk_recommendation.docx')

        # Example usage:
        # Assuming risk_recommendation is the generated recommendation
        risk_recommendation = (risk_recommendation)

        # Call the function to save to Word file
        save_to_word(risk_recommendation)

        print("Risk recommendation saved to 'risk_recommendation.docx'")
        
        
        # Function to send email with attachment
#         def send_email(to_email, subject, body, attachment_path):
#         # Email configuration
#             smtp_server = 'smtp.gmail.com'
#             smtp_port = 587  # Change it based on your SMTP server settings
#             smtp_username = 'propertyrisk2@gmail.com'
#             smtp_password = 'Chatgpt@123#'

#         # Create the email message
#             msg = MIMEMultipart()
#             msg['From'] = smtp_username
#             msg['To'] = to_email
#             msg['Subject'] = subject

#         # Attach body text
#             msg.attach(MIMEText(body, 'plain'))

#         # Attach the Word document
#             with open(attachment_path, 'rb') as file:
#                 attach = MIMEApplication(file.read(), _subtype="docx")
#                 attach.add_header('Content-Disposition', 'attachment', filename='risk_recommendation.docx')
#                 msg.attach(attach)

#         # Connect to the SMTP server
                
#             server = smtplib.SMTP(smtp_server, smtp_port)
#             server.starttls()
#             server.login(smtp_username, smtp_password)

#         # Send the email
#             server.sendmail(smtp_username, to_email, msg.as_string())

#         # Disconnect from the server
#             server.quit()

#         # Example usage:
        
#         to_email = 'mohit.uttam@icicilombard.com'
#         subject = 'Risk Recommendation Report'
#         body = 'Please find the risk recommendation report attached.'

#         print("Check 404!")
        
#         # Assuming 'risk_recommendation.docx' is the generated Word document
#         attachment_path = 'risk_recommendation.docx'

#         # Call the function to send the email
#         send_email(to_email, subject, body, attachment_path)

#         print(f"Email sent to {to_email} with the risk recommendation report.")
        
        print("Done???")
        
#         message = "Thanks for providing the information!" + "/n" + "Please check your email for detailed recommendation report"
#         return render_template('response.html', message=message)

#         result = "Thanks for providing the information!" + "/n" + "Please check your email for detatiled recommendation report"
        
#         return render_template('result.html', result=result)

    # Render the form page for GET requests
#     return render_template('form_copy.html')
#     return render_template('new_test.html')
    return render_template('new.html')

if __name__ == '__main__':
    app.run()


# In[ ]:




